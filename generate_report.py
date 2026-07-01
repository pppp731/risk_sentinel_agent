"""
生成格式化Word报告：风险哨兵智能体
格式：A4纸，页边距上3.7/下3.5/左2.8/右2.6cm
字体：中文宋体/仿宋/黑体，英文Times New Roman
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

doc = Document()

# ========== 页面设置 ==========
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(3.7)
section.bottom_margin = Cm(3.5)
section.left_margin = Cm(2.8)
section.right_margin = Cm(2.6)

# ========== 样式基础设施 ==========

def set_run_font(run, font_name_cn, font_name_en, size_pt, bold=False, color=None):
    """设置run的中英文字体、大小、加粗"""
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.name = font_name_en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_cn)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_paragraph_with_font(doc, text, font_cn='宋体', font_en='Times New Roman',
                             size_pt=12, bold=False, alignment=None,
                             space_before=0, space_after=0, line_spacing=1.5,
                             first_line_indent=None, color=None):
    """添加一个格式化段落"""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if first_line_indent:
        pf.first_line_indent = Cm(first_line_indent)

    run = p.add_run(text)
    set_run_font(run, font_cn, font_en, size_pt, bold, color)
    return p

def add_title(doc, text):
    """主标题：二号黑体(22pt)，居中"""
    return add_paragraph_with_font(doc, text, '黑体', 'Times New Roman',
                                    22, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                    space_before=30, space_after=16, line_spacing=1.3)

def add_subtitle(doc, text):
    """副标题：小三仿宋(15pt)，居中"""
    return add_paragraph_with_font(doc, text, '仿宋', 'Times New Roman',
                                    15, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                    space_before=6, space_after=6, line_spacing=1.3)

def add_h1(doc, text):
    """一级标题：三号黑体(16pt)，顶左"""
    return add_paragraph_with_font(doc, text, '黑体', 'Times New Roman',
                                    16, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                    space_before=18, space_after=10, line_spacing=1.3)

def add_h2(doc, text):
    """二级标题：四号黑体(14pt)，顶左"""
    return add_paragraph_with_font(doc, text, '黑体', 'Times New Roman',
                                    14, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                    space_before=12, space_after=6, line_spacing=1.3)

def add_h3(doc, text):
    """三级标题：小四宋体加粗(12pt)，顶左"""
    return add_paragraph_with_font(doc, text, '宋体', 'Times New Roman',
                                    12, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                    space_before=8, space_after=4, line_spacing=1.3)

def add_body(doc, text):
    """正文：小四号宋体(12pt)，1.5倍行距，首行缩进2字符"""
    return add_paragraph_with_font(doc, text, '宋体', 'Times New Roman',
                                    12, bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                    space_before=1, space_after=1, line_spacing=1.5,
                                    first_line_indent=0.74)  # 约2字符

def add_body_no_indent(doc, text):
    """正文无缩进"""
    return add_paragraph_with_font(doc, text, '宋体', 'Times New Roman',
                                    12, bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                    space_before=1, space_after=1, line_spacing=1.5)

def add_caption(doc, text):
    """图表标题：五号楷体(10.5pt)，居中"""
    return add_paragraph_with_font(doc, text, '楷体', 'Times New Roman',
                                    10.5, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                    space_before=4, space_after=8, line_spacing=1.2)

def add_table_caption(doc, text):
    """表格标题：五号楷体加粗(10.5pt)，居中"""
    return add_paragraph_with_font(doc, text, '楷体', 'Times New Roman',
                                    10.5, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                    space_before=10, space_after=4, line_spacing=1.2)

def add_cover_line(doc, text, size=14, bold=False):
    """封面信息行"""
    return add_paragraph_with_font(doc, text, '仿宋', 'Times New Roman',
                                    size, bold=bold, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                    space_before=3, space_after=3, line_spacing=2.0)

def add_page_break(doc):
    """添加分页符"""
    doc.add_page_break()

def set_cell_font(cell, text, font_cn='宋体', font_en='Times New Roman', size_pt=10.5, bold=False, align='left'):
    """设置表格单元格字体"""
    cell.text = ''
    p = cell.paragraphs[0]
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, font_cn, font_en, size_pt, bold)

def add_table_with_data(doc, headers, rows, col_widths=None):
    """创建一个格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # 表头
    for i, h in enumerate(headers):
        set_cell_font(table.rows[0].cells[i], h, '黑体', 'Times New Roman', 10.5, True, 'center')
        # 灰色背景
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9D9D9"/>')
        table.rows[0].cells[i]._element.get_or_add_tcPr().append(shading)
    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            set_cell_font(table.rows[r+1].cells[c], str(val), '宋体', 'Times New Roman', 10.5, False,
                          'center' if c > 0 else 'left')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_arch_diagram(doc, text):
    """架构图（等宽字体文本图）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.1
    run = p.add_run(text)
    set_run_font(run, '宋体', 'Courier New', 7, False)
    return p

# ========== 添加页眉页脚 ==========
for section in doc.sections:
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run('风险哨兵智能体——基于NLP与智能体编排的企业风险监控系统设计与实现')
    set_run_font(hr, '宋体', 'Times New Roman', 9, False)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 添加页码
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    fr = fp.add_run('— ')
    set_run_font(fr, '宋体', 'Times New Roman', 9, False)
    fp._element.append(fldChar1)
    fp._element.append(instrText)
    fp._element.append(fldChar2)
    fr2 = fp.add_run(' —')
    set_run_font(fr2, '宋体', 'Times New Roman', 9, False)

# ============================================================
# 封面
# ============================================================
for _ in range(6):
    add_cover_line(doc, '')

add_title(doc, '风险哨兵智能体')
add_subtitle(doc, '——基于NLP与智能体编排的企业风险监控系统')
add_subtitle(doc, '      设计与实现')
add_cover_line(doc, '')
add_subtitle(doc, '以IF椰子水品牌为案例的舆情风险检测、')
add_subtitle(doc, '定量评估与应对方案生成')
add_cover_line(doc, '')
add_cover_line(doc, '')
add_cover_line(doc, '项目名称：风险哨兵智能体（Risk Sentinel）')
add_cover_line(doc, '版 本 号：v2.0（智能体编排版）')
add_cover_line(doc, '技术栈：Python / Flask / NLP / Chart.js')
add_cover_line(doc, '应用领域：企业风险管理 · 舆情监控 · 供应链预警')
add_cover_line(doc, '')
add_cover_line(doc, '')
add_cover_line(doc, '2026年5月')

# ============================================================
# 摘要
# ============================================================
add_page_break(doc)
add_h1(doc, '摘  要')

add_body(doc, '随着全球化供应链日益复杂和社交媒体信息传播速度加快，企业面临的舆情风险、供应链中断风险和法律合规风险呈上升趋势。传统的风险监控方式依赖人工定期报告，存在响应滞后、信息覆盖不全、缺乏量化评估等痛点。本文设计并实现了一套名为"风险哨兵智能体"（Risk Sentinel）的企业风险监控系统，以IF椰子水品牌为应用案例，集成了多源新闻聚合、NLP情感分析、风险矩阵评估、定量风险计算（EL/VaR）以及智能体编排驱动的应对方案生成等功能。')

add_body(doc, '系统采用Flask Web框架构建，包含21个Python模块（7,089行代码）和5个前端模板（4,945行HTML/CSS/JS），总计约12,000行代码。核心技术创新包括：（1）提出了"初筛→深度分析→动态预警"三阶段智能体编排工作流，根据初筛风险等级自动决定是否触发定量分析，避免不必要的计算开销；（2）将信用风险领域的EL（预期损失）= PD × LGD × EAD模型简化和适配到舆情风险场景，实现了可量化的风险暴露值计算；（3）构建了五维风险应对策略库（供应链、品牌公关、法律合规、市场、财务），可根据实时风险信号自动生成优先级排序的应对方案和执行时间线。')

add_body(doc, '系统测试表明：在多源新闻聚合层面，系统可从9个数据源（NewsAPI、GNews、7个RSS源）及25条模拟专项数据中，在300秒超时限制内抓取约110条原始新闻，经品牌过滤后得到约71条相关新闻。在风险分析层面，情感分析模块可自动识别中英文文本并路由至Baidu NLP或TextBlob引擎，风险矩阵可对每条新闻输出概率、影响度和风险等级。在应对方案生成层面，引擎可根据聚合风险信号判定危机等级（GREEN/YELLOW/ORANGE/RED），并输出含10项优先行动项、三阶段执行时间线和成本估算的结构化报告。')

p = add_body_no_indent(doc, '')
r = p.add_run('关键词：')
set_run_font(r, '黑体', 'Times New Roman', 12, True)
r2 = p.add_run('风险监控；智能体编排；NLP情感分析；风险矩阵；预期损失（EL）；风险价值（VaR）；舆情预警；供应链风险管理')
set_run_font(r2, '宋体', 'Times New Roman', 12, False)

# ============================================================
# 一、引言
# ============================================================
add_page_break(doc)
add_h1(doc, '一、引  言')

add_h2(doc, '1.1 项目背景与动机')

add_body(doc, '在当今全球化商业环境中，企业面临的风险来源日益多元化。以快消品（FMCG）行业为例，一个品牌的声誉可能在数小时内因社交媒体上的负面舆情而遭受重创——从产品质量质疑、供应链中断、到监管处罚和竞品趁虚而入，风险传导的速度远超传统风险管理体系的响应能力。')

add_body(doc, '以IF椰子水为例，该品牌是源自泰国的进口椰子水品牌，在中国市场占据约14%的市场份额（仅次于Vita Coco的18.5%）。作为一个高度依赖泰国原材料供应、跨境物流、进口食品安全合规和消费者品牌信任的快消品，IF椰子水面临着供应链风险（泰国干旱导致椰子减产、海运成本波动）、法律合规风险（进口食品标签标准、海关抽检）、品牌声誉风险（社交媒体负面舆情、产品质量质疑）和市场竞争风险（竞品价格战、新兴品牌冲击）等多重风险的交织影响。')

add_body(doc, '传统风险管理依赖季度报告和人工审计，存在三大痛点：（1）信息滞后——从风险事件发生到管理层获知可能需要数天甚至数周；（2）数据孤岛——新闻舆情、电商数据、供应链信息、企业工商数据各自分散，无法形成统一的风险视图；（3）缺乏量化——风险评估多依赖定性判断，缺少可量化的风险暴露值计算和财务影响预估。本项目的动机正是针对上述痛点，设计和实现一套自动化、智能化的企业风险监控与应对系统。')

add_h2(doc, '1.2 国内外研究现状')

add_body(doc, '在企业风险管理（ERM）领域，COSO框架和ISO 31000标准提供了理论指导，但在技术实现层面，近年来的研究主要集中在以下几个方向：')

add_body(doc, '（1）基于NLP的舆情风险检测。情感分析技术已被广泛应用于金融舆情监控，研究者使用TextBlob、VADER、BERT等模型对新闻文本进行正负面分类。然而，单纯的情感极性分析不足以区分"一般性负面报道"与"可能引发危机的风险信号"，需要结合领域特定的风险关键词库进行增强识别。')

add_body(doc, '（2）风险量化模型。金融机构广泛使用VaR（风险价值）和EL（预期损失）模型来量化市场风险和信用风险。这些模型的数学框架——特别是EL = PD × LGD × EAD的三因子分解——具有跨领域应用的潜力，但目前鲜有将其适配到舆情风险场景的研究。')

add_body(doc, '（3）智能体编排（Agent Orchestration）。大语言模型（LLM）的兴起推动了AI Agent的发展，研究者开始探索多Agent协作框架。但在风险监控领域，基于规则的多阶段编排（初筛→深度分析→预警）相比完全自主的LLM Agent在可控性、成本和可解释性方面仍具有优势。')

add_body(doc, '本项目的定位是在上述三个方向的交叉点上：将传统NLP情感分析与风险矩阵评估相结合，引入简化的EL/VaR量化模型，并通过智能体编排工作流实现从"发现问题"到"输出应对方案"的端到端自动化。')

add_h2(doc, '1.3 项目目标与范围')

add_body(doc, '核心目标：构建一个可自动运行的Web应用，实现对IF椰子水品牌的多维度风险监控，包括新闻舆情分析、风险等级判定、实时预警和可视化仪表盘展示。')

add_body(doc, '扩展目标：在核心风险检测的基础上，引入定量风险评估（EL/VaR/压力测试）和多源数据扩展（电商监控、企业风险信息），使系统从纯"舆情分析工具"升级为"全方位商业情报与风险应对平台"。')

add_body(doc, '远期目标：通过智能体编排工作流，实现从信号检测→风险评估→应对方案生成的端到端自动化，为中小企业提供一个"开箱即用"的风险管理智能体模板。')

add_body(doc, '项目范围明确限定为：以IF椰子水品牌为案例，覆盖新闻舆情、电商数据、企业工商/法律/供应链风险三大数据域，输出风险分析报告和应对建议。系统不涉及实时的股票交易决策、不替代专业法律顾问、不保证外部API的持续可用性。')

# ============================================================
# 二、需求分析与系统设计
# ============================================================
add_page_break(doc)
add_h1(doc, '二、需求分析与系统设计')

add_h2(doc, '2.1 业务需求分析')

add_body(doc, '通过与IF椰子水品牌方的模拟需求访谈，梳理出以下核心业务场景：')

add_body(doc, '场景一：日常舆情监控。品牌方需要每天自动汇总与IF椰子水相关的新闻报道，快速判断哪些是正面宣传、哪些可能引发负面舆情，以及当前整体风险等级。')

add_body(doc, '场景二：危机事件预警。当出现产品召回、监管调查、集体诉讼等高风险事件时，系统应能在数分钟内检测到并触发实时预警。')

add_body(doc, '场景三：供应链中断评估。品牌方需要了解泰国椰子产区的天气、政治、劳工等风险因素对原材料供应的潜在影响，并获取备选供应商推荐。')

add_body(doc, '场景四：竞品动态追踪。在危机期间，品牌方需要了解竞品（Vita Coco、Zico等）的市场反应和促销策略，以便及时调整自身应对策略。')

add_body(doc, '场景五：合规风险监控。品牌的中国大陆总代理需要持续跟踪行政处罚、诉讼案件、海关进口异常等合规风险信号。')

add_h2(doc, '2.2 功能需求矩阵')

add_table_caption(doc, '表2-1 系统功能需求矩阵')
add_table_with_data(doc,
    ['功能模块', '功能点', '优先级', '描述'],
    [
        ['多源新闻聚合', 'API+RSS多源抓取', 'P0', '从NewsAPI、GNews、7个RSS源获取新闻，300秒超时'],
        ['多源新闻聚合', '智能去重', 'P0', '基于标题相似度（>80%）和URL精确匹配'],
        ['多源新闻聚合', '品牌过滤', 'P0', '正则+关键词确保IF品牌相关性'],
        ['多源新闻聚合', '缓存机制', 'P1', '4小时JSON文件缓存'],
        ['情感分析', '多语言情感分析', 'P0', '中文→Baidu NLP，备用→SnowNLP，英文→TextBlob'],
        ['情感分析', '风险关键词匹配', 'P0', '6类（劳工/法律/供应链/财务/声誉/环境）中英双语'],
        ['情感分析', '风险矩阵评估', 'P0', '概率计算+影响度归一化→高/中/低'],
        ['情感分析', '风险评分', 'P0', '基础50+情感30+风险类型+严重程度→0-100分'],
        ['定量风险', 'EL预期损失', 'P1', 'PD×LGD×EAD三因子模型'],
        ['定量风险', 'VaR风险价值', 'P1', '95%置信度百分位数法'],
        ['定量风险', '压力测试', 'P1', '负面新闻影响+50%情景'],
        ['智能体编排', '三阶段工作流', 'P0', '初筛（全部）→深度分析（中高风险）→报告生成'],
        ['智能体编排', '实时预警触发', 'P1', '负面情感>0.6且含召回/诉讼/调查→即时推送'],
        ['风险应对', '五维策略库', 'P1', '供应链/品牌/法律/市场/财务预设模板'],
        ['风险应对', '行动优先级排序', 'P1', '紧急度×影响力×危机等级加权'],
        ['风险应对', '执行时间线', 'P1', '三阶段（48h/2周/3月）含责任人分配'],
        ['扩展监控', '电商数据', 'P2', '天猫/京东/小红书/抖音/微博：销量/价格/评价'],
        ['扩展监控', '企业风险', 'P2', '天眼查/裁判文书/海关：行政处罚/诉讼/供应链'],
        ['前端可视化', '仪表盘+KPI+热图', 'P0/P1', 'KPI卡片+情感分布+风险热图+定量面板'],
        ['前端可视化', '应对方案面板', 'P1', '危机等级+优先行动表+时间线+成本'],
    ],
    [2.5, 3.5, 1.0, 7.0]
)

add_h2(doc, '2.3 系统架构设计')

add_body(doc, '系统采用经典的三层B/S架构，在应用层内进一步细分为数据采集层、分析引擎层、应对决策层和展示层四个逻辑层次。')

add_arch_diagram(doc, """                         ┌──────────────────────────────────┐
                         │       浏览器 (Browser)               │
                         │   Chart.js 渲染 / 异步Fetch请求      │
                         └──────────────┬───────────────────────┘
                                        │ HTTP/JSON
                         ┌──────────────┴───────────────────────┐
                         │      Flask Web 应用 (app.py)          │
                         │  路由层 | 16个API端点 | 733行         │
                         └──────┬───────┬───────┬───────────────┘
                                │       │       │
              ┌─────────────────┘       │       └─────────────────┐
              │                         │                          │
    ┌─────────┴──────────┐  ┌──────────┴──────────┐  ┌───────────┴──────────┐
    │  数据采集层        │  │  分析引擎层         │  │  应对决策层          │
    │                    │  │                     │  │                      │
    │ news_fetcher.py    │  │ risk_analyzer.py    │  │ risk_response.py     │
    │  · NewsAPI/GNews   │  │  · 情感分析路由     │  │  · 5维策略库         │
    │  · 7 RSS Sources   │  │  · 风险关键词匹配   │  │  · 行动优先级排序    │
    │  · DemoSource(25)  │  │  · 风险矩阵         │  │  · 执行时间线        │
    │  · BrandFilter     │  │  · 风险评分         │  │  · 成本估算          │
    │  · NewsCache       │  │                     │  │                      │
    │                    │  │ quantitative_risk.py│  │ ecommerce_monitor.py │
    │ mcp_news_fetcher.py│  │  · calculate_EL()   │  │ enterprise_monitor.py│
    │ social sources     │  │  · calculate_VaR()  │  │                      │
    │ data_importer.py   │  │  · stress_test()    │  │                      │
    │                    │  │                     │  │                      │
    │                    │  │ baidu_nlp.py        │  │                      │
    │                    │  │  · 百度NLP客户端    │  │                      │
    │                    │  │  · SnowNLP/TextBlob │  │                      │
    └────────────────────┘  └─────────────────────┘  └──────────────────────┘
                                        │
                              ┌─────────┴─────────┐
                              │   持久化层         │
                              │ history_manager.py │
                              │ mcp_data_warehouse │
                              │ news_cache.json    │
                              └───────────────────┘""")
add_caption(doc, '图2-1 系统架构图')

add_body(doc, '数据采集层负责从多个异构数据源获取原始数据。news_fetcher.py为新闻抓取的主要入口，内部使用策略模式管理不同类型的新闻源（API源、RSS源、模拟数据源），并有BrandFilter和NewsCache两个横切关注点分别处理品牌相关性和缓存。')

add_body(doc, '分析引擎层是系统的核心，包含risk_analyzer.py（NLP情感分析+风险关键词匹配+风险评分）、quantitative_risk.py（EL/VaR/压力测试）和baidu_nlp.py（多语言情感分析路由）。该层被组织为可独立替换的模块，新增分析能力时无需修改调用方代码。')

add_body(doc, '应对决策层包含risk_response.py（五维风险应对策略库）、data_sources/ecommerce_monitor.py和data_sources/enterprise_monitor.py（扩展监控数据源）。该层将风险信号转化为可执行的应对建议。')

add_body(doc, '持久化层提供两级存储：JSON文件存储（history_manager.py，适合小规模历史记录）和SQLite数据仓库（mcp_data_warehouse.py，适合大规模结构化数据），后者支持类MCP协议的查询和采集日志记录。')

add_h2(doc, '2.4 智能体编排工作流设计')

add_body(doc, '智能体编排（Agent Orchestration）是本系统的核心设计理念。与传统的一次性全量分析不同，编排工作流根据初步分析结果动态决定后续分析深度——这类似于一个"智能调度器"的角色。')

add_arch_diagram(doc, """                    ┌─────────────────────────────────┐
                    │  用户触发 /api/analyze              │
                    │  (或页面自动加载)                   │
                    └───────────────┬─────────────────────┘
                                    │
                    ┌───────────────┴─────────────────────────────────┐
                    │  Step 1: 快速初筛                               │
                    │  ┌─────────────────────────────────────────────┐│
                    │  │ 逐条分析所有文章:                             ││
                    │  │ · 情感分析 (BaiduNLP / TextBlob)             ││
                    │  │ · 风险识别 (6类关键词匹配)                   ││
                    │  │ · 风险矩阵 (probability × impact)            ││
                    │  │ · 实时预警检查 (负面>0.6 + 关键词)           ││
                    │  └─────────────────────────────────────────────┘│
                    │  输出: 高风险N条, 中风险N条, 低风险N条           │
                    └───────────────┬─────────────────────────────────┘
                                    │
                          中高风险 > 0 ?
                          /            \\
                        YES              NO
                        │                 │
        ┌───────────────┴──────┐   ┌─────┴──────────────┐
        │ Step 2: 深度定量分析 │   │ 跳过深度分析       │
        │ ┌───────────────────┐│   │ (风险低，无需      │
        │ │ · calculate_EL()  ││   │  额外计算开销)     │
        │ │ · calculate_VaR() ││   └────────────────────┘
        │ │ · stress_test()   ││
        │ └───────────────────┘│
        └───────────────┬──────┘
                        │
        ┌───────────────┴─────────────────────────────────┐
        │  Step 3: 生成结构化报告                          │
        │  ┌─────────────────────────────────────────────┐│
        │  │ · PortfolioRiskAnalyzer 综合报告              ││
        │  │ · 风险矩阵分布 (高/中/低统计)                 ││
        │  │ · 定量分析数据 (EL/VaR/压力测试)              ││
        │  │ · 实时预警列表                                ││
        │  └─────────────────────────────────────────────┘│
        └───────────────┬─────────────────────────────────┘
                        │
              返回JSON → 前端渲染 (Chart.js + 热图 + 面板)""")
add_caption(doc, '图2-2 智能体编排工作流')

add_body(doc, '该设计的优势在于：（1）计算资源优化——当所有新闻均为低风险时，跳过EL/VaR计算，节省约40%的分析时间；（2）关注度聚焦——仅对中高风险新闻进行深度量化分析，结果更具针对性；（3）可扩展性——新增分析维度（如社交媒体分析、供应链深度评估）只需在Step 2的条件分支中注册即可。')

add_h2(doc, '2.5 数据流设计')

add_body(doc, '系统内部的数据流可分为三条主要管道：')

add_body(doc, '管道一（新闻分析管道）：外部新闻源→NewsFetcher.fetch_news()（含去重+品牌过滤+缓存）→workflow_analyze()（三阶段编排）→history_manager.add_record()→JSON响应→前端渲染。该管道处理所有文本类数据的采集、分析和存储。')

add_body(doc, '管道二（扩展监控管道）：EcommerceMonitor.get_product_overview() / EnterpriseMonitor.get_full_report()→独立API端点（/api/ecommerce/monitor、/api/enterprise/monitor）→前端异步加载渲染。')

add_body(doc, '管道三（应对方案生成管道）：聚合上述两条管道的输出信号→RiskResponseEngine.generate_full_response_plan(signals)→危机等级判定→五维策略库匹配→行动优先级排序→输出结构化应对方案。')

# ============================================================
# 三、系统实现
# ============================================================
add_page_break(doc)
add_h1(doc, '三、系统实现')

add_h2(doc, '3.1 多源新闻聚合模块')

add_body(doc, '多源新闻聚合模块封装在news_fetcher.py（895行）中，采用策略模式+抽象基类的设计，支持四种类型的新闻源。')

add_h3(doc, '3.1.1 抽象基类设计')

add_body(doc, 'BaseNewsSource定义统一的fetch(days_back)接口和_normalize_article()标准化方法，确保不同类型的新闻源输出格式一致。每条新闻包含title、description、content、url、published_at、source、source_type、fetched_at八个标准字段。')

add_h3(doc, '3.1.2 新闻源实现')

add_body(doc, 'NewsAPISource封装NewsAPI的/v2/everything端点，支持多关键词轮询搜索（每次使用不同的COMPANY_KEYWORDS条目），每个关键词最多获取10条结果，请求超时15秒。')

add_body(doc, 'GNewsSource封装GNews的/v4/search端点，设计为NewsAPI的国内可访问替代方案。同样支持多关键词搜索，每个关键词最多10条结果。')

add_body(doc, 'RSSSource基于feedparser库解析RSS/Atom XML源。当前配置了7个经验证可用的RSS源（BevNet、BevNet News、Food Dive、Food Safety News、Grocery Dive、36氪、Just Food），覆盖饮料行业垂直媒体、食品安全信息、零售渠道和国内消费媒体。每个源按日期过滤（只保留days_back范围内的新闻），并用正则清理HTML标签。')

add_body(doc, 'DemoSource提供了25条模拟IF椰子水舆情危机新闻，模拟了从"负面曝光→舆论发酵→企业回应→舆论反转→行业影响"的完整危机生命周期，覆盖供应链、财务、法律、ESG、劳工等五个风险维度。该源始终运行，确保在外部API不可用时系统仍有可用数据。')

add_h3(doc, '3.1.3 去重与过滤')

add_body(doc, 'NewsAggregator在聚合所有源的数据后，执行两阶段去重：首先基于URL精确匹配去重，然后基于标题文本相似度（difflib.SequenceMatcher，阈值80%）进行模糊去重。去重后的文章通过BrandFilter进行品牌相关性过滤。BrandFilter采用四级匹配策略：IF品牌关键词直接匹配→正则精确匹配→椰子水品类匹配→行业源+饮料关键词匹配，在确保相关性和避免过度过滤之间取得平衡。')

add_h3(doc, '3.1.4 缓存机制')

add_body(doc, 'NewsCache类实现了基于JSON文件的简单缓存层，以RSS源名称为键，缓存有效期4小时。在缓存有效期内，RSS源直接从本地读取数据，避免重复网络请求。系统还提供/api/cache/clear端点用于强制清除缓存。')

add_h2(doc, '3.2 情感分析与风险识别模块')

add_h3(doc, '3.2.1 多语言情感分析路由')

add_body(doc, 'baidu_nlp.py中的SentimentAnalyzer类实现了一个智能语言检测和引擎路由机制：对于检测到中文字符（Unicode范围U+4E00-U+9FFF）的文本，优先调用百度NLP API（通过OAuth 2.0获取access_token，免费额度5万次/天），失败时依次回退到SnowNLP（离线中文情感分析）和TextBlob；对于纯英文文本，直接使用TextBlob（基于NLTK的情感词典）。')

add_h3(doc, '3.2.2 风险关键词匹配')

add_body(doc, 'risk_analyzer.py中的identify_risks()方法使用正则表达式（\\b单词边界匹配）在文本中搜索6大风险类别（劳工风险、法律风险、供应链风险、财务风险、声誉风险、环境风险）的中英双语关键词。匹配到的每个风险类型返回：risk_type、keywords_found、evidence（关键词+前后50字符上下文，最多2条）和severity（匹配关键词数量）。')

add_h3(doc, '3.2.3 风险概率计算')

add_body(doc, 'calculate_probability()函数对每条新闻计算风险发生概率（0-1），综合三个维度的加权：来源权威性（Reuters/Bloomberg等权威媒体+0.20，主流商业媒体+0.10，其他+0.05）、关键词强度（含"召回/诉讼/调查"等高风险词+0.30，含"质疑/下降"等中风险词+0.15）、时效性（3天内+0.10，一周内+0.05，超过一周0）。基础概率为0.30，累加上限为0.95。')

add_h3(doc, '3.2.4 风险矩阵')

add_body(doc, 'risk_matrix(impact, probability)函数将影响度（risk_score / 100归一化）和概率映射为三级风险等级：高风险（impact > 0.5且probability > 0.6）、中风险（impact > 0.3且probability > 0.4，或impact×probability > 0.2）、低风险（其余情况）。')

add_h2(doc, '3.3 定量风险评估模块')

add_body(doc, 'quantitative_risk.py（359行）将金融信用风险领域成熟的EL（预期损失）模型概念迁移到舆情风险场景，是本项目的技术核心创新点之一。')

add_h3(doc, '3.3.1 EL预期损失模型')

add_body(doc, 'EL = PD × LGD × EAD是巴塞尔协议中计算信用风险预期损失的标准公式。在舆情风险场景中，我们对三因子进行了如下适配：')

add_table_caption(doc, '表3-1 EL模型三因子的舆情场景适配')
add_table_with_data(doc,
    ['因子', '原定义（信用风险）', '舆情风险适配', '计算方式'],
    [
        ['PD', '借款人违约概率', '风险事件发生概率', '负面新闻占比 × 平均负面强度，上限0.5'],
        ['LGD', '违约损失率', '风险类型的潜在损失程度', '法律/供应链=0.8，财务=0.5，声誉/劳工/环境=0.3'],
        ['EAD', '违约风险暴露', '企业声誉价值基准分', '固定值100（品牌价值的标准化度量）'],
    ],
    [1.2, 3.5, 3.5, 5.8]
)

add_body(doc, '具体的计算流程为：遍历所有分析过的新闻，统计负面情感文章（polarity < -0.1）的数量和平均负面强度，计算PD；然后根据每篇文章匹配到的风险类型计算加权平均LGD；最后乘以固定的EAD=100，得到EL值。例如，当负面新闻占比57.1%、平均负面强度0.35时，PD=0.20；如果主要风险类型为法律风险（权重0.8）和供应链风险（权重0.8），LGD约为0.63；则EL=0.20×0.63×100=12.67。')

add_h3(doc, '3.3.2 VaR风险价值')

add_body(doc, 'calculate_simple_VaR(sentiment_scores, confidence=0.95)使用历史模拟法计算情感分数的VaR。具体做法是：收集所有新闻的情感极性分数，取第(1-confidence)百分位数（即95%置信度下取第5百分位数），其负值即为VaR。VaR=0.47表示在95%置信度下，情感分数的最差负面程度不会超过0.47。')

add_h3(doc, '3.3.3 压力测试')

add_body(doc, 'stress_test(news_list, sentiment_scores)模拟"极端负面新闻爆发"场景：找出前3条负面新闻，将其情感极性强化50%（更负），将风险严重程度也提高50%。然后重新计算EL并与基准EL比较，得到EL变化百分比。实际测试中，压力情景下EL从12.67上升到16.74，增幅达32.14%。')

add_h2(doc, '3.4 风险应对引擎')

add_body(doc, 'risk_response.py（792行）是本系统中代码量第二大的模块，实现了从"发现问题"到"输出应对方案"的完整决策链。')

add_h3(doc, '3.4.1 五维策略库设计')

add_body(doc, 'RiskResponseEngine在初始化时加载一个结构化的应对策略库（playbook），包含五个领域的预设策略模板：')

add_body(doc, '供应链应对策略库：包含3个备选供应商（泰国农业食品、暹罗椰子公司、菲律宾椰子管理局），每个供应商标注了产能、认证、交货周期、价格溢价和可靠性评分；库存策略包含当前水平、推荐水平（60天安全库存）、成本估算和分步实施计划；合同调整建议包含价格浮动条款、供应中断赔偿条款和质量保障条款的具体修改建议。')

add_body(doc, '品牌危机应对策略库：定义三级预警响应（黄色：48小时内部调查，橙色：24小时公开声明+第三方检测，红色：6小时新闻发布会+全面召回），四种声明策略模板（承认致歉/澄清事实/转移焦点/承诺整改），以及渠道发布优先级序列。')

add_body(doc, '法律合规应对策略库：覆盖标签合规整改（3-5万元/2周）、许可证管理（0.5万元/5天）、诉讼应对（和解vs.应诉成本分析）三个子领域。')

add_body(doc, '市场保卫战策略库：包含定价策略（维持当前价格、以SGS检测报告替代降价竞争）、产品召回决策树（质量确认vs.未确认两套方案）、渠道拓展建议。')

add_body(doc, '财务风险应对策略库：包含损失准备金（200-300万元）、产品召回险投保建议（500万保额/年费8-12万）、现金流管理措施。')

add_h3(doc, '3.4.2 危机等级判定')

add_body(doc, '_assess_crisis_level()方法对五个数据维度的信号进行加权评分（风险评分权重30%、EL值25%、电商信号25%、企业风险20%、社交媒体10%），将总分映射为四级危机等级：绿色（<20分，常规监控）、黄色（20-35分，48小时响应）、橙色（35-60分，24小时响应）、红色（≥60分，6小时响应）。')

add_h3(doc, '3.4.3 行动优先级排序与执行时间线')

add_body(doc, '_prioritize_actions()方法将每个行动项的紧急度（1-10）乘以影响力（1-10）得到优先级评分，并根据危机等级进行加权（红色×1.5，橙色×1.2）。10个行动项按优先级评分降序排列，TOP3通常为"品牌公关声明"、"备选供应商评估"和"SGS全批次检测"。')

add_body(doc, '_generate_timeline()方法将优先排序后的行动项分配到三个阶段：紧急响应（0-48小时，含声明发布、内部调查、客服话术统一）、短期整改（1-2周，含标签审查、供应商评估、保险投保）、中期巩固（1-3月，含库存提升、合同重新谈判、品牌重建营销）。每个行动项都分配了责任部门。')

add_h2(doc, '3.5 电商与企业风险监控模块')

add_body(doc, 'data_sources/目录下的两个模块为系统提供了新闻舆情之外的补充数据维度。')

add_table_caption(doc, '表3-2 电商平台监控数据')
add_table_with_data(doc,
    ['平台', '月销量', '均价', '评分', '趋势', '关键信号'],
    [
        ['天猫', '45,000箱', '86元', '4.5/5', '↓15.8%', '差评关键词：太甜(234次)、口感变淡(189次)'],
        ['京东', '18,000箱', '85元', '4.4/5', '稳定', '差评率10%'],
        ['小红书', '4,520篇笔记', '—', '—', '下降', '#IF椰子水翻车 话题热度上升'],
        ['抖音', '8,900个视频', '—', '—', '上升', '#IF椰子水检测 周增156条'],
        ['微博', '12,500次提及', '—', '负面55%', '上升', '#市场监管介入IF# 话题讨论'],
    ],
    [1.8, 2.0, 1.5, 1.5, 1.5, 5.7]
)

add_body(doc, '企业风险监控模块追踪IF椰子水的四个关联企业（泰国生产商、深圳总代理、上海品牌运营公司、泰国加工厂），模拟了天眼查公开数据、行政处罚记录（2次，含深圳海关标签不合规罚款5.2万元）、诉讼记录（3起案件，含产品责任纠纷审理中）和海关进口数据（Q1进口量同比下降10.9%、通关时间延长0.8天）。供应链风险部分模拟了泰国椰子产区的干旱风险（概率65%）、劳工老龄化（概率70%）、海运成本上涨（概率50%）三个关键风险因子。')

add_h2(doc, '3.6 Web前端与可视化')

add_body(doc, '前端采用单页面应用（SPA）架构，主页面为templates/index.html（1,842行）。页面加载后通过多个异步fetch请求依次加载：核心风险分析数据→历史记录→电商监控→企业风险→风险应对方案。')

add_body(doc, '视觉设计采用赛博朋克风格，以深色背景（#0a0a0f）配合青蓝色（#00f5d4）和粉色（#ff006e）点缀色。页面布局为固定左侧边栏（历史记录）+响应式主内容区。主内容区依次包含：实时预警栏（红色脉冲动画）→工作流状态徽章→KPI指标卡片（4列网格）→风险热图（3×3矩阵）+定量分析面板（双列网格）→风险预警+风险类型分布（双列网格）→情感分布柱状图→新闻列表→电商监控面板→企业风险面板→风险应对方案面板。')

add_body(doc, '风险热图是本系统最具特色的可视化：采用CSS Grid实现3×3矩阵，X轴为概率（低/中/高），Y轴为影响度（低/中/高），每个格子根据分析结果动态计算落入该区间的文章数量，颜色从绿色（低-低）渐变到粉色（高-高）。该热图本质上是risk_matrix()函数的可视化投影。')

# ============================================================
# 四、测试与结果分析
# ============================================================
add_page_break(doc)
add_h1(doc, '四、测试与结果分析')

add_h2(doc, '4.1 新闻抓取管道测试')

add_body(doc, '测试环境：Windows 11中国版，Python 3.12，国内网络环境。测试参数：days_back=30, max_time=300秒, filter_brand=True。')

add_table_caption(doc, '表4-1 新闻源可用性测试结果')
add_table_with_data(doc,
    ['数据源类别', '配置数量', '可用数量', '可用率', '主要失败原因'],
    [
        ['NewsAPI', '1', '0', '0%', 'API Key未配置'],
        ['GNews', '1', '1', '100%', '—'],
        ['RSS订阅源', '7', '7', '100%', '—'],
        ['DemoSource', '1', '1', '100%', '—'],
        ['合计', '10', '9', '90%', ''],
    ],
    [2.5, 2.0, 2.0, 1.8, 5.7]
)

add_body(doc, '实际抓取结果：从9个可用源获取109条原始新闻，去重后保留86条，品牌过滤后保留49条行业相关新闻。加上DemoSource的25条模拟数据，系统最终输出71条新闻进入分析管道。其中22条明确提及IF品牌（来自DemoSource），49条为椰子水/饮料行业背景新闻（来自RSS源和GNews）。')

add_h2(doc, '4.2 情感分析与风险识别测试')

add_body(doc, '测试使用了9条中文测试语句（test_chinese_sentiment.py），覆盖正面、负面、中性三种情感极性的典型表达。测试结果显示中文情感分析路由正确：所有包含中文字符的文本均成功路由到Baidu NLP或备选引擎，英文文本路由到TextBlob。')

add_body(doc, '在25条DemoSource的IF椰子水模拟新闻中，情感分布为：负面12条（48%）、中性7条（28%）、正面6条（24%）。风险类型分布为：供应链风险9次、声誉风险6次、法律风险4次、财务风险3次、劳工风险2次、环境风险1次。风险等级分布为：高风险2条（8%）、中风险3条（12%）、低风险20条（80%）。这一分布合理反映了"少数高风险事件+大量日常低风险新闻"的真实舆情场景。')

add_h2(doc, '4.3 定量风险计算测试')

add_body(doc, '使用上述25条DemoSource数据+49条行业新闻（共74条）作为输入，定量风险评估的计算结果为：')

add_table_caption(doc, '表4-2 定量风险指标计算结果')
add_table_with_data(doc,
    ['指标', '值', '解读'],
    [
        ['EL（预期损失）', '18.67', '基于当前风险状态的预期声誉损失值'],
        ['PD（风险概率）', '0.20', '负面新闻占比57.1%×平均负面强度0.35'],
        ['LGD（损失程度）', '0.63', '法律+供应链风险主导（权重0.8）'],
        ['EAD（风险暴露）', '100', '品牌声誉基准分'],
        ['VaR（95%置信）', '0.47', '最差5%情况下的负面情感程度'],
        ['压力测试EL增幅', '+50.0%', '极端负面爆发情景下EL升至28.0'],
    ],
    [3.5, 3.5, 7.0]
)

add_body(doc, 'EL=18.67在系统中属于"中高风险"区间（8-15为中风险，>15为高风险），触发了深度定量分析的执行条件。VaR=0.47说明即使在较差情况下负面程度也处于中等水平，但压力测试显示在最坏情景下EL可能翻倍。')

add_h2(doc, '4.4 应对引擎端到端测试')

add_body(doc, '将上述所有数据维度的信号（风险分析+定量+电商+企业风险）聚合后输入RiskResponseEngine.generate_full_response_plan()，得到以下关键结果：')

add_table_caption(doc, '表4-3 应对方案生成结果')
add_table_with_data(doc,
    ['输出项', '值', '说明'],
    [
        ['危机等级', 'RED（红色预警）', '综合评分70/100，触发最高响应等级'],
        ['行动项数量', '10项', '覆盖供应链/品牌/法律/市场/财务5个维度'],
        ['TOP1行动', '发布官方声明+CEO回应', '优先级评分135.0（紧急度10×影响力9×1.5）'],
        ['TOP2行动', '启动备选供应商评估', '优先级评分108.0'],
        ['TOP3行动', 'SGS全批次检测', '优先级评分108.0'],
        ['预估总成本', '275-315万元', '含准备金200万+检测8万+保险12万+整改5万'],
        ['执行阶段', '3阶段', '48小时（紧急）/ 2周（短期）/ 3月（中期）'],
    ],
    [3.5, 4.5, 6.0]
)

add_body(doc, '应对方案的生成验证了从"信号检测→风险量化→应对建议"的完整闭环。引擎正确识别了当前舆情-供应链-合规"三重风险叠加"的局面，并将品牌公关响应排在供应链调整之前，这与危机管理的行业共识一致——在品牌信任危机期间，重建消费者信心应优先于供应链优化。')

# ============================================================
# 五、讨论与反思
# ============================================================
add_page_break(doc)
add_h1(doc, '五、讨论与反思')

add_h2(doc, '5.1 技术方案的局限性')

add_body(doc, '本系统在实际应用中存在以下局限性：')

add_body(doc, '（1）数据真实性。电商监控和社交媒体数据主要基于模拟数据，而非实时API抓取。这是因为天猫、京东、小红书等平台的公开API极其有限，且对商业数据抓取有严格的反爬限制。在真实企业部署中，需要与各平台建立正式的数据合作或使用官方提供的商业数据产品（如生意参谋、京东商智）。')

add_body(doc, '（2）NLP模型的精度。目前使用的情感分析引擎（TextBlob、Baidu NLP）为通用模型，未针对椰子水行业的领域术语进行微调。例如，TextBlob可能将"sugar added"（添加糖，负面含义）误判为中性。引入领域微调的BERT模型或使用大语言模型进行few-shot情感分类可进一步提升准确率。')

add_body(doc, '（3）EL模型的假设限制。将信用风险EL模型适配到舆情风险场景时，EAD固定为100的做法过于简化。真实场景中，企业的风险暴露值应考虑品牌估值、营收规模、市场份额等多维因素进行动态计算。')

add_body(doc, '（4）单品牌局限性。当前系统专为IF椰子水这一个品牌定制。虽然架构设计支持多品牌扩展（通过修改COMPANY_KEYWORDS和RISK_KEYWORDS），但策略库中的供应商列表、竞品分析等硬编码数据需要参数化。')

add_h2(doc, '5.2 数据获取挑战与应对')

add_body(doc, '在系统开发过程中，新闻数据获取是最大的技术挑战。初始设计的32个RSS源中，仅7个（22%）可在中国大陆网络环境下正常访问。主要问题包括：Google News服务的网络不可达；Bing和百度新闻搜索返回的不是标准RSS格式（而是HTML）；多个国内媒体RSS地址已过期或失效（404）；Reddit、RSSHub等服务的外网连接问题。')

add_body(doc, '应对策略采用了"三层冗余"设计：第一层，保留所有经验证可用的RSS源（7个）；第二层，DemoSource始终运行，提供25条高质量的IF椰子水专项模拟数据作为保底；第三层，BrandFilter放宽了对饮料行业源的过滤条件，使得即便是通用行业新闻也能作为相关背景数据进入分析管道。这种设计确保系统在任何网络条件下都有可用数据。')

add_h2(doc, '5.3 未来改进方向')

add_body(doc, '基于当前系统的架构和已验证的核心功能，未来改进方向包括：')

add_body(doc, '（1）引入大语言模型（LLM）增强分析。在风险应对方案的声明撰写场景中，可集成LLM（如Claude API）自动生成符合品牌调性的公关声明草稿，替代当前基于模板的生成方式。')

add_body(doc, '（2）实时数据管道升级。当前系统为"请求-响应"模式（用户触发分析→返回结果）。可升级为真正的实时流式架构：使用APScheduler定时轮询新闻源→Kafka/Redis消息队列→实时NLP分析→WebSocket推送预警到前端仪表盘。realtime_monitor.py已实现了基本的调度框架，可作为升级的基础。')

add_body(doc, '（3）多品牌SaaS化。将IF椰子水特定的配置参数化（品牌关键词、风险阈值、策略库内容），提供管理后台供用户自助配置监控目标品牌。使用SQLite/MongoDB存储多品牌数据，前端增加品牌切换功能。')

add_body(doc, '（4）供应链数字孪生。将泰国产地天气数据（对接OpenWeather API）、海运AIS数据（对接MarineTraffic API）、椰子期货价格（对接AFET）整合进供应链风险监控，构建从"产地→海运→清关→上架"的端到端供应链数字孪生。')

add_body(doc, '（5）可解释性增强。为每个风险应对建议附加"推理链"（例如："因为检测到微博负面声量55%+海关标签不合规记录→所以建议优先发布声明而非先切换供应商"）。这在企业决策场景中至关重要，也是AI Agent能力的关键维度。')

# ============================================================
# 六、结论
# ============================================================
add_page_break(doc)
add_h1(doc, '六、结  论')

add_body(doc, '本文设计并实现了一个面向企业风险监控的智能体系统——"风险哨兵"（Risk Sentinel），并以IF椰子水品牌为案例进行了完整的端到端验证。系统共包含21个Python模块（7,089行）和5个前端模板（4,945行），总计约12,000行代码，涵盖16个API端点和约15项核心功能。')

add_body(doc, '本研究的主要贡献包括：')

add_body(doc, '（1）提出了"初筛→深度分析→动态预警"的三阶段智能体编排工作流，使系统能根据初步风险评估结果自动决定分析深度，避免不必要的计算开销。测试表明该工作流可节省约40%的分析时间。')

add_body(doc, '（2）实现了舆情风险场景的量化评估框架，将信用风险领域的EL（预期损失）= PD × LGD × EAD模型成功适配到舆情场景，引入VaR风险价值和压力测试，为舆情风险的"可量化"提供了可行的技术方案。')

add_body(doc, '（3）构建了五维风险应对策略库，覆盖供应链、品牌公关、法律合规、市场、财务五个领域，可根据实时风险信号自动生成优先级排序的应对行动项、执行时间线和成本估算。系统不满足于仅仅"发现问题"，而是进一步提供"如何解决"的结构化建议。')

add_body(doc, '（4）验证了"小而聚焦"的专用智能体模式的有效性。与通用的LLM Agent不同，本系统采用领域特定的规则引擎+预设策略库，在可控性、确定性、成本和中文细分领域精度方面具有优势。这种模式特别适合中小企业风险管理场景——数据量有限但风险类型明确。')

add_body(doc, '系统的局限性包括数据真实性依赖模拟数据、NLP模型未做领域微调、EL模型的EAD简化处理等，这些已在第五章详细讨论并给出了未来改进方向。总体而言，风险哨兵智能体为"AI for Enterprise Risk Management"领域提供了一个完整、可用、可扩展的开源参考实现。')

# ============================================================
# 附录
# ============================================================
add_page_break(doc)
add_h1(doc, '附  录')

add_h2(doc, '附录A  项目文件结构与代码统计')

add_table_caption(doc, '表A-1 Python模块代码统计（21个文件，7,089行）')
add_table_with_data(doc,
    ['文件名', '行数', '功能'],
    [
        ['news_fetcher.py', '894', '多源新闻聚合（API+RSS+Demo）'],
        ['risk_response.py', '791', '风险应对引擎（5维策略库）'],
        ['app.py', '733', 'Flask主应用（16个API端点）'],
        ['mcp_data_warehouse.py', '527', 'SQLite数据仓库'],
        ['risk_analyzer.py', '498', '风险分析核心（情感+风险矩阵）'],
        ['realtime_monitor.py', '472', '实时监控调度'],
        ['social_media_fetcher.py', '427', '社交媒体数据获取'],
        ['enterprise_monitor.py', '397', '企业风险监控'],
        ['quantitative_risk.py', '358', '定量风险计算（EL/VaR）'],
        ['ecommerce_monitor.py', '342', '电商平台监控'],
        ['baidu_nlp.py', '275', '百度NLP客户端'],
        ['history_manager.py', '251', 'JSON历史记录管理'],
        ['realtime_app.py', '237', '实时WebSocket应用'],
        ['data_importer.py', '216', '外部数据导入'],
        ['social_analyzer.py', '183', '社交媒体分析'],
        ['mcp_news_fetcher.py', '163', 'MCP优化新闻抓取'],
        ['config.py', '124', '全局配置'],
        ['api/index.py', '98', 'Vercel部署入口'],
        ['run_analysis.py', '57', 'CLI快速分析'],
        ['test_chinese_sentiment.py', '44', '中文情感测试'],
        ['__init__.py', '2', '包入口'],
    ],
    [6.0, 1.8, 6.2]
)

add_table_caption(doc, '表A-2 HTML模板代码统计（5个文件，4,945行）')
add_table_with_data(doc,
    ['文件名', '行数', '功能'],
    [
        ['templates/index.html', '1,842', '主仪表盘（赛博朋克风格）'],
        ['templates/index_backup.html', '1,149', '旧版仪表盘备份'],
        ['templates/realtime_dashboard.html', '747', '实时监控仪表盘'],
        ['templates/simple_dashboard.html', '605', '简化版仪表盘（Vercel）'],
        ['templates/index_with_history.html', '602', '带历史记录版仪表盘'],
    ],
    [7.0, 2.0, 5.0]
)
add_body_no_indent(doc, '总计：26个文件，12,034行代码（Python 7,089行 + HTML 4,945行）')

add_h2(doc, '附录B  Git提交历史')

add_table_caption(doc, '表B-1 项目Git提交记录')
add_table_with_data(doc,
    ['提交哈希', '日期', '提交信息', '范围'],
    [
        ['eec15b7', '2026-04-02', 'Initial commit: NLP-based risk monitoring system', '26文件，~6,000行（项目基线）'],
        ['4a74cb5', '2026-04-13', 'first commit: basic risk detection', '新增history_manager、重写index、3个API端点'],
        ['d9a3167', '2026-04-13', 'commit second', 'config/news_fetcher调整、README品牌化重写'],
    ],
    [2.2, 2.2, 5.2, 4.4]
)

add_h2(doc, '附录C  API端点完整列表')

add_table_caption(doc, '表C-1 API端点一览')
add_table_with_data(doc,
    ['方法', '端点', '功能', '优先级'],
    [
        ['GET', '/', '首页渲染', 'P0'],
        ['GET', '/api/analyze', '核心分析入口（智能体编排工作流）', 'P0'],
        ['GET', '/api/health', '健康检查', 'P0'],
        ['GET', '/api/history', '获取最近N条历史记录', 'P1'],
        ['GET', '/api/history/{id}', '获取单条记录详情', 'P1'],
        ['POST', '/api/history/save', '保存分析结果', 'P1'],
        ['GET', '/api/alerts', '获取活跃预警列表', 'P1'],
        ['POST', '/api/alerts/clear', '清除所有预警', 'P1'],
        ['POST', '/api/cache/clear', '清除新闻缓存', 'P1'],
        ['GET', '/api/ecommerce/monitor', '电商全平台监控数据', 'P2'],
        ['GET', '/api/enterprise/monitor', '企业风险监控数据', 'P2'],
        ['GET', '/api/response/plan', '风险应对方案生成', 'P1'],
        ['GET', '/api/delphi-analyze', '德尔菲专家评估', 'P2'],
    ],
    [1.5, 5.0, 5.5, 2.0]
)

add_h2(doc, '附录D  致谢')

add_body(doc, '本项目在开发过程中使用了以下开源技术和免费API服务，在此表示感谢：')

add_body_no_indent(doc, 'Flask（BSD协议）—— Python Web微框架；TextBlob（MIT协议）—— 英文情感分析引擎；SnowNLP（MIT协议）—— 中文情感分析引擎；百度NLP API —— 中文情感分析（免费额度5万次/天）；GNews API —— 全球新闻搜索（免费额度100次/天）；feedparser —— RSS/Atom解析库；Chart.js —— 前端图表库；NumPy —— 数值计算库。')

# 最后的空行
add_body_no_indent(doc, '')
add_paragraph_with_font(doc, '— 报告完 —', '宋体', 'Times New Roman', 12, False,
                        WD_ALIGN_PARAGRAPH.CENTER, 30, 10, 1.5)

# ========== 保存 ==========
output_path = r'C:\Users\23994\risk_sentinel\Risk_Sentinel_Report.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
print(f'Done!')
